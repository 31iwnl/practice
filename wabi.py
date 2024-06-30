import cv2
import numpy as np
import os
from matplotlib import pyplot as plt
from docx import Document
from docx.shared import Inches

def check_image_quality(image):
    # Проверка на размытость изображения
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    laplacian_var = cv2.Laplacian(gray, cv2.CV_64F).var()
    if laplacian_var < 100:
        blur_quality = "Плохое качество (размытость)"
    else:
        blur_quality = "Хорошее качество"

    # Определение разрешения изображения
    height, width = image.shape[:2]
    if width >= 1920 and height >= 1080:
        resolution_quality = "Высокое качество (1080p или выше)"
    elif width >= 1280 and height >= 720:
        resolution_quality = "Среднее качество (720p)"
    elif width >= 640 and height >= 360:
        resolution_quality = "Низкое качество (360p)"
    else:
        resolution_quality = "Очень низкое качество (ниже 360p)"

    return {
        "Размытость": blur_quality,
        "Разрешение": f"{width}x{height}",
        "Качество разрешения": resolution_quality
    }

def convert_to_monochrome(image):
    # Преобразование изображения в монохромный вид (черно-белый)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    _, monochrome = cv2.threshold(gray, 127, 255, cv2.THRESH_BINARY)
    return monochrome

def detect_orientation(image):
    # Определение ориентации изображения
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 100, 200)
    lines = cv2.HoughLines(edges, 1, np.pi / 180, 200)
    if lines is not None:
        angles = [np.degrees(theta) for rho, theta in lines[:, 0]]
        average_angle = np.mean(angles)
        if 45 < average_angle < 135:
            orientation = "Горизонтальная ориентация"
        else:
            orientation = "Вертикальная ориентация"
    else:
        orientation = "Не удалось определить ориентацию"
    return orientation

def correct_orientation(image):
    # Определение и коррекция ориентации изображения
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 250, 200)
    lines = cv2.HoughLines(edges, 1, np.pi / 180, 200)
    if lines is not None:
        angles = [np.degrees(theta) for rho, theta in lines[:, 0]]
        median_angle = np.median(angles)
        # Корректируем угол для правильного поворота
        angle_to_rotate = (median_angle - 90) if median_angle > 45 else median_angle
        # Поворачиваем изображение на нужный угол
        (h, w) = image.shape[:2]
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, angle_to_rotate, 1.0)
        rotated = cv2.warpAffine(image, M, (w, h) , borderValue=(255,255,255))
        return rotated
    else:
        return image

def save_image(image, filename, folder="output_images"):
    # Создаем папку, если она не существует
    if not os.path.exists(folder):
        os.makedirs(folder)
    # Сохраняем изображение
    cv2.imwrite(os.path.join(folder, filename), image)

def save_images_to_docx(images, docx_path):
    # Создаем документ Word
    doc = Document()
    for image_path, title in images:
        doc.add_heading(title, level=1)
        doc.add_picture(image_path, width=Inches(6))
        doc.add_page_break()
    doc.save(docx_path)

# Загрузка изображения
image_path = 'vov1.jpg'
image = cv2.imread(image_path)

# Определение качества изображения
quality = check_image_quality(image)
print("Качество изображения:")
print("Размытость:", quality["Размытость"])
print("Разрешение:", quality["Разрешение"])
print("Качество разрешения:", quality["Качество разрешения"])

# Преобразование изображения в монохромный вид
monochrome_image = convert_to_monochrome(image)

# Определение ориентации изображения
orientation = detect_orientation(image)
print("Ориентация изображения:", orientation)

# Коррекция ориентации изображения
corrected_image = correct_orientation(image)

# Определение новой ориентации изображения
new_orientation = detect_orientation(corrected_image)
print("Новая ориентация изображения:", new_orientation)

# Сохранение изображений
save_image(image, "original_image.jpg")
save_image(monochrome_image, "monochrome_image.jpg")
save_image(corrected_image, "corrected_image.jpg")

# Сохранение изображений в Word документ
images = [

    ("output_images/corrected_image.jpg", "")
]
save_images_to_docx(images, "docs/test/test.docx")

# Отображение изображений
plt.figure(figsize=(15, 5))
plt.subplot(1, 3, 1)
plt.title('Оригинальное изображение')
plt.imshow(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))

plt.subplot(1, 3, 2)
plt.title('Монохромное изображение')
plt.imshow(monochrome_image, cmap='gray')

plt.subplot(1, 3, 3)
plt.title('Корректированное изображение')
plt.imshow(cv2.cvtColor(corrected_image, cv2.COLOR_BGR2RGB))

plt.show()
